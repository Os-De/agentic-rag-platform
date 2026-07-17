"""File → text extraction (Phase 1: txt/md/pdf/docx/html)."""

import io
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".html", ".htm"}


class UnsupportedFileType(ValueError):
    pass


class EmptyDocument(ValueError):
    pass


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:  # keep tabular knowledge, row per line
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Unsupported file type '{suffix}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    if suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    elif suffix in {".html", ".htm"}:
        text = _extract_html(data)
    else:
        text = data.decode("utf-8", errors="replace")

    if not text.strip():
        raise EmptyDocument(f"No extractable text in '{filename}' (scanned PDF? needs OCR)")
    return text

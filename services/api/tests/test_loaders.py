import io

import pytest

from app.ingestion.loaders import EmptyDocument, UnsupportedFileType, extract_text


def test_markdown_passthrough():
    assert "hello" in extract_text("a.md", b"# hello world")


def test_html_strips_tags_and_scripts():
    html = b"<html><script>evil()</script><body><h1>Title</h1><p>Body text</p></body></html>"
    text = extract_text("page.html", html)
    assert "Title" in text and "Body text" in text
    assert "evil" not in text


def test_docx_extracts_paragraphs_and_tables():
    import docx

    document = docx.Document()
    document.add_paragraph("Phase 4 is about security.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "role"
    table.rows[0].cells[1].text = "admin"
    buf = io.BytesIO()
    document.save(buf)

    text = extract_text("doc.docx", buf.getvalue())
    assert "Phase 4 is about security." in text
    assert "role | admin" in text


def test_unsupported_extension():
    with pytest.raises(UnsupportedFileType):
        extract_text("virus.exe", b"nope")


def test_empty_document():
    with pytest.raises(EmptyDocument):
        extract_text("blank.txt", b"   \n  ")
